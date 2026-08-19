from pathlib import Path
import re
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
MD = ROOT / "docs" / "learn-claude-code-notes.md"
OUT = ROOT / "docx" / "QianAgent-Learn-Claude-Code-Notes.docx"

SANS = "Noto Sans CJK SC"
SERIF = "Noto Serif CJK SC"
MONO = "Noto Sans Mono CJK SC"


def set_run_font(run, name, size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def no_row_split(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(OxmlElement("w:cantSplit"))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("— ")
    set_run_font(run, SANS, 8.5, color=(110, 118, 129))
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)
    run2 = paragraph.add_run(" —")
    set_run_font(run2, SANS, 8.5, color=(110, 118, 129))


def add_inline_runs(paragraph, text, base_font=SERIF, base_size=10.5):
    link_pattern = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
    pos = 0
    expanded = ""
    for match in link_pattern.finditer(text):
        expanded += text[pos:match.start()] + f"{match.group(1)} ({match.group(2)})"
        pos = match.end()
    expanded += text[pos:]

    for part in re.split(r"(\*\*.*?\*\*|`[^`]+`)", expanded):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, base_font, base_size, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, MONO, 9.5, color=(35, 88, 140))
        else:
            run = paragraph.add_run(part)
            set_run_font(run, base_font, base_size)


def add_code_block(doc, lines):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade_cell(cell, "F6F8FA")
    set_cell_margins(cell, 110, 150, 110, 150)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    run = paragraph.add_run("\n".join(lines).rstrip())
    set_run_font(run, MONO, 8.5, color=(31, 35, 40))


def add_md_table(doc, rows):
    if len(rows) < 2:
        return
    data = [rows[0]] + rows[2:]
    cols = max(len(row) for row in data)
    table = doc.add_table(rows=0, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, source in enumerate(data):
        row = table.add_row()
        no_row_split(row)
        for col_index in range(cols):
            cell = row.cells[col_index]
            set_cell_margins(cell, 80, 95, 80, 95)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index == 0:
                shade_cell(cell, "EAF2F8")
            text = source[col_index] if col_index < len(source) else ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline_runs(paragraph, text, SANS if row_index == 0 else SERIF, 8.8)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True


def parse_table_row(line):
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = SERIF
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), SERIF)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.38
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 28, (20, 48, 78)),
        ("Subtitle", 12, (89, 99, 110)),
        ("Heading 1", 18, (24, 65, 100)),
        ("Heading 2", 15, (24, 65, 100)),
        ("Heading 3", 11.5, (36, 82, 120)),
    ]:
        style = styles[name]
        style.font.name = SANS
        style._element.rPr.rFonts.set(qn("w:eastAsia"), SANS)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor(*color)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
    styles["Heading 1"].paragraph_format.space_before = Pt(0)
    styles["Heading 1"].paragraph_format.space_after = Pt(10)
    styles["Heading 2"].paragraph_format.space_before = Pt(18)
    styles["Heading 2"].paragraph_format.space_after = Pt(7)
    styles["Heading 3"].paragraph_format.space_before = Pt(12)
    styles["Heading 3"].paragraph_format.space_after = Pt(5)


def add_cover(doc):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(74)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("QianAgent")
    set_run_font(run, SANS, 18, bold=True, color=(36, 82, 120))

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(8)
    run = paragraph.add_run("Learn Claude Code · 12 话 Agent Harness 结构化笔记")
    set_run_font(run, SANS, 25, bold=True, color=(20, 48, 78))

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(16)
    run = paragraph.add_run("从 Agent Loop 到 Tool Dispatch、Task DAG、Multi-Agent 与 Worktree Isolation")
    set_run_font(run, SERIF, 11.5, color=(89, 99, 110))

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(36)
    run = paragraph.add_run("整理版 · 2026-08")
    set_run_font(run, SANS, 9.5, color=(110, 118, 129))

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(28)
    run = paragraph.add_run("MODEL  →  HARNESS  →  ENVIRONMENT")
    set_run_font(run, MONO, 10, bold=True, color=(36, 82, 120))

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(60)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(
        "说明：本文件由系列转录材料重新组织，不是逐字稿。\n"
        "明显 ASR 误识别已统一；外部扩展内容会明确标注为 QianAgent 仓库映射。"
    )
    set_run_font(run, SERIF, 9, color=(89, 99, 110))
    doc.add_page_break()


def add_reading_map(doc):
    heading = doc.add_paragraph(style="Heading 1")
    heading.add_run("阅读地图")
    intro = doc.add_paragraph()
    add_inline_runs(
        intro,
        "这套材料可以按四层理解：先建立 Loop，再把动作变成 Tool；随后把计划、知识和上下文外部化；最后处理并发、多 Agent、自治与环境隔离。",
    )

    rows = [
        ("层次", "章节", "核心问题"),
        ("控制循环", "S01–S02", "模型如何持续行动，并安全连接真实环境"),
        ("状态与上下文", "S03–S07", "计划、知识、上下文与任务依赖如何显式化"),
        ("并发与协作", "S08–S10", "耗时任务与长期 Agent 如何并行、通信、协商"),
        ("自治与隔离", "S11–S12", "空闲 Agent 如何自己找活，以及并行修改如何互不污染"),
        ("系统合成", "最终章", "所有能力如何仍围绕同一个 Agent Loop 工作"),
    ]
    table = doc.add_table(rows=0, cols=3)
    for row_index, values in enumerate(rows):
        row = table.add_row()
        no_row_split(row)
        for col_index, text in enumerate(values):
            cell = row.cells[col_index]
            set_cell_margins(cell, 90, 100, 90, 100)
            if row_index == 0:
                shade_cell(cell, "EAF2F8")
            paragraph = cell.paragraphs[0]
            add_inline_runs(paragraph, text, SANS if row_index == 0 else SERIF, 9)
            if row_index == 0:
                for run in paragraph.runs:
                    run.bold = True

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(14)
    run = paragraph.add_run(
        "目录：0 Overview · 1 Agent Loop · 2 Tool Use · 3 TodoWrite · 4 Subagent · "
        "5 Skill Loading · 6 Context Compact · 7 Task System · 8 Background Tasks · "
        "9 Agent Teams · 10 Team Protocols · 11 Autonomous Agents · 12 Worktree Isolation · "
        "13 系统合成 · 14 概念对照 · 15 QianAgent 映射 · 16 完整任务流 · 附录"
    )
    set_run_font(run, SERIF, 9.2, color=(89, 99, 110))
    doc.add_page_break()


def add_header_footer(section):
    header = section.header.paragraphs[0]
    header.text = "QianAgent · Learn Claude Code 12 话结构化笔记"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        set_run_font(run, SANS, 8, color=(110, 118, 129))
    add_page_number(section.footer.paragraphs[0])


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.65)
    section.left_margin = Cm(1.85)
    section.right_margin = Cm(1.85)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.75)
    configure_styles(doc)
    add_header_footer(section)
    add_cover(doc)
    add_reading_map(doc)

    lines = MD.read_text(encoding="utf-8").splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("## 0."))
    lines = lines[start:]

    in_code = False
    code_lines = []
    table_rows = []
    current_h2_seen = False

    def flush_table():
        nonlocal table_rows
        if table_rows:
            add_md_table(doc, table_rows)
            table_rows = []

    for raw in lines:
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            if not in_code:
                in_code = True
                code_lines = []
            else:
                add_code_block(doc, code_lines)
                in_code = False
                code_lines = []
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            table_rows.append(parse_table_row(line))
            continue
        flush_table()

        if line.startswith("## "):
            if current_h2_seen:
                doc.add_page_break()
            current_h2_seen = True
            paragraph = doc.add_paragraph(style="Heading 1")
            add_inline_runs(paragraph, line[3:], SANS, 18)
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(style="Heading 2")
            add_inline_runs(paragraph, line[4:], SANS, 15)
            continue
        if line.startswith("#### "):
            paragraph = doc.add_paragraph(style="Heading 3")
            add_inline_runs(paragraph, line[5:], SANS, 11.5)
            continue
        if line.strip() == "---" or not line.strip():
            continue
        if line.startswith("> "):
            table = doc.add_table(rows=1, cols=1)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            cell = table.cell(0, 0)
            shade_cell(cell, "F3F6F8")
            set_cell_margins(cell, 100, 130, 100, 130)
            paragraph = cell.paragraphs[0]
            add_inline_runs(paragraph, line[2:], SERIF, 9.7)
            continue

        match = re.match(r"^[-*]\s+(.*)", line)
        if match:
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.25
            add_inline_runs(paragraph, match.group(1), SERIF, 10.2)
            continue

        match = re.match(r"^(\d+)\.\s+(.*)", line)
        if match:
            paragraph = doc.add_paragraph(style="List Number")
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.25
            add_inline_runs(paragraph, match.group(2), SERIF, 10.2)
            continue

        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Cm(0.74)
        paragraph.paragraph_format.line_spacing = 1.36
        paragraph.paragraph_format.space_after = Pt(5)
        add_inline_runs(paragraph, line, SERIF, 10.5)

    flush_table()
    if in_code and code_lines:
        add_code_block(doc, code_lines)

    props = doc.core_properties
    props.title = "QianAgent × Learn Claude Code：12 话 Agent Harness 结构化笔记"
    props.subject = "Agent Loop、Tool Dispatch、Task System、Multi-Agent 与 Worktree Isolation"
    props.author = "QianAgent Documentation"
    props.keywords = "QianAgent, Claude Code, Agent Harness, Agent Runtime, Tool Use, Multi-Agent"

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT, OUT.stat().st_size)


if __name__ == "__main__":
    build_docx()
